"""
Script to generate test resume files (PDF and DOCX).
Run this script to create sample resumes for testing.

Usage:
    python generate_test_files.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import os


def generate_test_docx():
    """Generate a realistic test resume in DOCX format."""
    doc = Document()

    # Name (Title)
    name = doc.add_paragraph('John Smith')
    name.runs[0].font.size = Pt(24)
    name.runs[0].bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact = doc.add_paragraph('john.smith@email.com | (555) 123-4567 | San Francisco, CA')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Professional Summary
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(
        'Experienced Software Engineer with 5+ years of expertise in full-stack development, '
        'specializing in Python, JavaScript, and cloud technologies. Proven track record of '
        'delivering scalable applications and leading development teams.'
    )

    # Skills
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(
        'Python, JavaScript, React, Node.js, FastAPI, Django, PostgreSQL, MongoDB, '
        'Docker, Kubernetes, AWS, Git, CI/CD, Agile, REST API, GraphQL, Machine Learning'
    )

    # Work Experience
    doc.add_heading('Work Experience', level=2)

    # Job 1
    job1_title = doc.add_paragraph('Senior Software Engineer')
    job1_title.runs[0].bold = True
    doc.add_paragraph('Tech Innovations Inc. | 2021 - Present')
    doc.add_paragraph(
        '• Led development of microservices architecture serving 1M+ users\n'
        '• Implemented CI/CD pipelines reducing deployment time by 60%\n'
        '• Mentored team of 5 junior developers\n'
        '• Technologies: Python, FastAPI, React, AWS, Docker, PostgreSQL'
    )

    # Job 2
    job2_title = doc.add_paragraph('Software Engineer')
    job2_title.runs[0].bold = True
    doc.add_paragraph('StartupCo | 2019 - 2021')
    doc.add_paragraph(
        '• Developed full-stack web applications using React and Node.js\n'
        '• Built RESTful APIs handling 10k+ requests per minute\n'
        '• Implemented automated testing increasing code coverage to 85%\n'
        '• Technologies: JavaScript, React, Node.js, Express, MongoDB'
    )

    # Education
    doc.add_heading('Education', level=2)
    edu = doc.add_paragraph('Bachelor of Science in Computer Science')
    edu.runs[0].bold = True
    doc.add_paragraph('University of California, Berkeley | 2019')

    # Certifications
    doc.add_heading('Certifications', level=2)
    doc.add_paragraph('• AWS Certified Solutions Architect')
    doc.add_paragraph('• Certified Kubernetes Administrator (CKA)')

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'test_data', 'test_resume.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Created test DOCX resume: {output_path}")


def generate_test_pdf():
    """Generate a realistic test resume in PDF format."""
    output_path = os.path.join(os.path.dirname(__file__), 'test_data', 'test_resume.pdf')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1e40af',
        spaceAfter=6,
        alignment=1  # Center
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='#1e40af',
        spaceAfter=6,
        spaceBefore=12
    )

    # Story (content)
    story = []

    # Name
    story.append(Paragraph('Jane Doe', title_style))
    story.append(Paragraph(
        'jane.doe@email.com | (555) 987-6543 | Austin, TX',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.3 * inch))

    # Professional Summary
    story.append(Paragraph('Professional Summary', heading_style))
    story.append(Paragraph(
        'Results-driven Full Stack Developer with 6+ years of experience building '
        'scalable web applications. Expert in modern JavaScript frameworks, Python backend '
        'development, and cloud infrastructure. Strong problem-solver with excellent '
        'communication skills.',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Skills
    story.append(Paragraph('Skills', heading_style))
    story.append(Paragraph(
        'Python, JavaScript, TypeScript, React, Angular, Vue, Django, Flask, FastAPI, '
        'Node.js, Express, SQL, PostgreSQL, MySQL, MongoDB, Redis, Docker, Kubernetes, '
        'AWS, Azure, GCP, Git, Jenkins, CI/CD, Agile, Scrum, TDD, REST API, GraphQL, '
        'Microservices',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Work Experience
    story.append(Paragraph('Work Experience', heading_style))

    story.append(Paragraph('<b>Lead Software Engineer</b>', styles['Normal']))
    story.append(Paragraph('CloudTech Solutions | 2020 - Present', styles['Normal']))
    story.append(Paragraph(
        '• Architected and deployed cloud-native applications on AWS serving 5M+ users<br/>'
        '• Led cross-functional team of 8 developers and designers<br/>'
        '• Reduced infrastructure costs by 40% through optimization<br/>'
        '• Technologies: Python, React, TypeScript, AWS, Kubernetes, PostgreSQL',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph('<b>Full Stack Developer</b>', styles['Normal']))
    story.append(Paragraph('WebDev Corp | 2018 - 2020', styles['Normal']))
    story.append(Paragraph(
        '• Developed e-commerce platform processing $2M+ in transactions monthly<br/>'
        '• Implemented real-time features using WebSockets<br/>'
        '• Improved application performance by 50% through optimization<br/>'
        '• Technologies: JavaScript, React, Node.js, MongoDB, Redis',
        styles['Normal']
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Education
    story.append(Paragraph('Education', heading_style))
    story.append(Paragraph('<b>M.S. in Computer Science</b>', styles['Normal']))
    story.append(Paragraph('Stanford University | 2018', styles['Normal']))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph('<b>B.S. in Software Engineering</b>', styles['Normal']))
    story.append(Paragraph('University of Texas at Austin | 2016', styles['Normal']))
    story.append(Spacer(1, 0.2 * inch))

    # Certifications
    story.append(Paragraph('Certifications', heading_style))
    story.append(Paragraph('• AWS Certified Developer - Associate', styles['Normal']))
    story.append(Paragraph('• Google Cloud Professional Cloud Architect', styles['Normal']))
    story.append(Paragraph('• Certified Scrum Master (CSM)', styles['Normal']))

    # Build PDF
    doc.build(story)
    print(f"✅ Created test PDF resume: {output_path}")


def main():
    """Generate all test files."""
    print("Generating test resume files...")
    generate_test_docx()
    generate_test_pdf()
    print("\n✅ All test files generated successfully!")
    print("\nFiles created:")
    print("  - backend/app/tests/test_data/test_resume.docx")
    print("  - backend/app/tests/test_data/test_resume.pdf")


if __name__ == '__main__':
    main()
